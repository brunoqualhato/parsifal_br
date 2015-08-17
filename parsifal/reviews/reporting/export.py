# coding: utf-8

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def export_review_to_docx(review):
    document = Document()

    h = document.add_heading(review.title, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph('')

    authors = list()
    authors.append(review.author.profile.get_screen_name())
    for author in review.co_authors.all():
        authors.append(author.profile.get_screen_name())
    p = document.add_paragraph(', '.join(authors))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph('')

    if review.description:
        document.add_paragraph(review.description)

    document.add_heading('1 Protocol', level=2)

    if review.objective:
        document.add_paragraph(review.objective)

    '''
        PICOC
    '''
    document.add_heading('1.1 PICOC', level=3)

    p = document.add_paragraph('', style='List Bullet')
    p.add_run('Population: ').bold = True
    p.add_run(review.population)

    p = document.add_paragraph('', style='List Bullet')
    p.add_run('Intervention: ').bold = True
    p.add_run(review.intervention)

    p = document.add_paragraph('', style='List Bullet')
    p.add_run('Comparison: ').bold = True
    p.add_run(review.comparison)

    p = document.add_paragraph('', style='List Bullet')
    p.add_run('Outcome: ').bold = True
    p.add_run(review.outcome)

    p = document.add_paragraph('', style='List Bullet')
    p.add_run('Context: ').bold = True
    p.add_run(review.context)

    '''
        Research Questions
    '''
    document.add_heading('1.2 Research Questions', level=3)

    for question in review.research_questions.all():
        document.add_paragraph(question.question, style='List Number')

    '''
        Keywords and Synonym
    '''
    document.add_heading('1.3 Keywords and Synonyms', level=3)

    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Keyword'
    hdr_cells[1].text = 'Synonyms'

    for keyword in review.get_keywords():
        row_cells = table.add_row().cells
        row_cells[0].text = keyword.description
        row_cells[1].text = ', '.join(keyword.synonyms.all().values_list('description', flat=True))

    '''
        Search String
    '''
    document.add_heading('1.4 Search String', level=3)
    document.add_paragraph(review.get_generic_search_string().search_string)

    '''
        Sources
    '''
    document.add_heading('1.5 Sources', level=3)

    for source in review.sources.all():
        text = source.name
        if source.url:
            text = u'{0} ({1})'.format(source.name, source.url)
        document.add_paragraph(text, style='List Bullet')

    '''
        Selection Criteria
    '''
    document.add_heading('1.6 Selection Criteria', level=3)

    p = document.add_paragraph()
    p.add_run('Inclusion Criteria:').bold = True
    for criteria in review.get_inclusion_criterias():
        document.add_paragraph(criteria.description, style='List Bullet')

    p = document.add_paragraph()
    p.add_run('Exclusion Criteria:').bold = True
    for criteria in review.get_exclusion_criterias():
        document.add_paragraph(criteria.description, style='List Bullet')

    return document